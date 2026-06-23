"""Generate a configurable preliminary datasheet DOCX."""

from __future__ import annotations

import argparse
import json
import sys
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path
from xml.etree import ElementTree as ET


REL_NS = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


SAMPLE_CONFIG = {
    "product": {
        "name": "<target_part_number>",
        "subtitle": "<datasheet subtitle>",
        "revision": "<revision or TBD>",
        "date": "<YYYY/MM/DD or TBD>",
        "package": "<package or TBD>",
        "compatibility": "<pin-to-pin compatibility statement or TBD>",
    },
    "sources": {
        "input_documents": [
            {
                "role": "style_template",
                "path": "<format template .docx preferred, .pdf visual reference allowed>",
                "purpose": "format/style/layout only; do not preserve placeholder body text",
            },
            {
                "role": "reference_datasheet",
                "path": "<reference or competitor datasheet .pdf/.docx>",
                "purpose": "content framework, comparison points, and source-risk discovery",
            },
            {
                "role": "company_prior_datasheet",
                "path": "<prior company datasheet .docx/.pdf or omit when unavailable>",
                "purpose": "company terminology, confirmed reusable facts, and reusable assets",
            },
            {
                "role": "target_product_facts",
                "path": "<authoritative target product facts .docx/.pdf or omit when unavailable>",
                "purpose": "target product claims and specifications",
            },
        ],
        "analysis_report": "<analysis.json>",
        "source_map": "<source_map.json>",
        "asset_manifest": "<assets.json>",
    },
    "template_guidance": [
        {
            "source_role": "style_template",
            "source_ref": "<template.docx comment id or source_map id>",
            "instruction": "<template comment or annotation guidance>",
            "decision": "<applied|not_applicable|needs_clarification>",
            "evidence": "<where the guidance was applied, or why it does not apply>",
        }
    ],
    "risk_markers": {
        "missing": "TBD - NEED NCS CONFIRMATION",
        "weaker": "COMPETITOR BETTER - NEED REVIEW",
        "competitor_source": "SOURCE FROM COMPETITOR - NEED NCS CONFIRMATION",
    },
    "output_options": {
        "include_audit_sections": False,
    },
    "sections": [
        {
            "title": "<section title from source map>",
            "blocks": [
                {
                    "type": "paragraph",
                    "text": "<agent-authored paragraph from confirmed target facts>",
                    "source_ref": "<source_map fact id>",
                    "confidence": "confirmed",
                },
                {
                    "type": "paragraph",
                    "risk": "competitor_source",
                    "text": "<reference-derived wording that still needs NCS confirmation>",
                    "source_ref": "<source_map fact id>",
                    "confidence": "needs_confirmation",
                },
            ],
        },
        {
            "title": "<bulleted section title from source map>",
            "blocks": [
                {
                    "type": "bullets",
                    "items": [
                        {"text": "<confirmed feature>", "source_ref": "<source_map fact id>", "confidence": "confirmed"},
                        {
                            "risk": "missing",
                            "text": "<missing or unresolved feature>",
                            "source_ref": "<source_map fact id or TBD>",
                            "confidence": "needs_confirmation",
                        },
                    ],
                }
            ],
        },
        {
            "title": "<table section title from source map>",
            "blocks": [
                {
                    "type": "table",
                    "headers": ["<column 1>", "<column 2>", "<status/source column>"],
                    "rows": [["<value>", "<value>", "TBD - NEED NCS CONFIRMATION"]],
                    "source_ref": "<source_map fact id>",
                }
            ],
        },
        {
            "title": "<revision/history section title if required>",
            "blocks": [
                {
                    "type": "table",
                    "headers": ["<revision column>", "<date column>", "<description column>"],
                    "rows": [["<revision or TBD>", "<YYYY/MM/DD or TBD>", "<draft description>"]],
                }
            ],
        },
    ],
    "confirmed_claims": [],
    "risk_items": [
        {"kind": "missing", "text": "<project-specific unresolved item>"},
        {"kind": "weaker", "text": "<project-specific reference advantage needing review>"},
    ],
}


def load_config(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def init_config(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(SAMPLE_CONFIG, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def add_run(paragraph, text: str, *, bold: bool = False, red: bool = False):
    from docx.shared import RGBColor

    run = paragraph.add_run(text)
    run.bold = bold
    if red:
        run.font.color.rgb = RGBColor(192, 0, 0)
    return run


def add_heading(doc, text: str):
    paragraph = doc.add_paragraph()
    try:
        paragraph.style = "DS标题1"
    except Exception:
        try:
            paragraph.style = "Heading 1"
        except Exception:
            pass
    add_run(paragraph, text.upper(), bold=True)


def add_title(doc, text: str):
    """Document title that tolerates templates lacking the built-in Title style.

    Falls back to a manually sized bold paragraph (22pt) when the cloned
    template has no 'Title' style, mirroring the fallback pattern used by
    add_heading for custom-heading templates.
    """
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    styled = False
    try:
        paragraph.style = "Title"
        styled = True
    except Exception:
        styled = False
    run = paragraph.add_run(text)
    run.bold = True
    if not styled:
        run.font.size = Pt(22)
    return paragraph


def add_warning(doc, marker: str, text: str):
    paragraph = doc.add_paragraph()
    add_run(paragraph, f"{marker}: {text}", bold=True, red=True)


def placeholder_png(label: str) -> Path:
    try:
        import warnings

        with warnings.catch_warnings():
            warnings.simplefilter("ignore", DeprecationWarning)
            import fitz  # type: ignore
    except Exception as exc:
        raise RuntimeError("PyMuPDF is required to create image placeholders. Install PyMuPDF.") from exc

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    tmp_path = Path(tmp.name)
    tmp.close()
    document = fitz.open()
    page = document.new_page(width=360, height=150)
    page.draw_rect(fitz.Rect(1, 1, 359, 149), color=(0.2, 0.45, 0.75), fill=(0.9, 0.94, 0.98), width=1)
    page.insert_textbox(
        fitz.Rect(18, 42, 342, 110),
        label[:220],
        fontsize=12,
        fontname="helv",
        color=(0.1, 0.25, 0.45),
        align=1,
    )
    page.get_pixmap(alpha=False).save(tmp_path)
    document.close()
    return tmp_path


def add_placeholder_image(doc, marker: str, label: str) -> None:
    from docx.shared import Inches

    image_path = placeholder_png(label)
    try:
        doc.add_picture(str(image_path), width=Inches(2.8))
    finally:
        image_path.unlink(missing_ok=True)
    add_warning(doc, marker, label)


def add_risk_paragraph(doc, marker: str, text: str, *, style=None):
    paragraph = doc.add_paragraph(style=style)
    add_run(paragraph, f"{marker}: ", bold=True, red=True)
    add_run(paragraph, text)
    return paragraph


def risk_marker(markers: dict, risk: str | None) -> str | None:
    if not risk:
        return None
    return markers.get(risk, risk)


def add_table(doc, headers: list[str], rows: list[list[str]], red_markers=None):
    """Generic table; cells whose text is a known risk marker render red+bold.

    Keeps risk markers visually prominent even inside tables (pin status,
    revision notes), matching the skill's requirement that all risk markers
    be醒目 (prominent) and verify_datasheet.py's --risk-marker red check.
    """
    from docx.shared import RGBColor

    red_markers = set(red_markers or [])
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.text = ""
            run = paragraph.add_run(value)
            if value in red_markers:
                run.bold = True
                run.font.color.rgb = RGBColor(192, 0, 0)
    doc.add_paragraph()


def normalize_table_rows(headers: list[str], rows: list) -> list[list[str]]:
    normalized = []
    for row in rows:
        if isinstance(row, dict):
            normalized.append([str(row.get(header, "")) for header in headers])
        else:
            normalized.append([str(value) for value in row])
    return normalized


def add_bullet(doc, text: str, marker: str | None = None):
    try:
        paragraph = doc.add_paragraph(style="List Bullet")
    except Exception:
        paragraph = doc.add_paragraph()
    if marker:
        add_run(paragraph, f"{marker}: ", bold=True, red=True)
    add_run(paragraph, text)


def render_block(doc, block: dict, markers: dict, marker_values: set[str]) -> None:
    block_type = block.get("type", "paragraph")
    marker = risk_marker(markers, block.get("risk"))

    if block_type == "paragraph":
        text = str(block.get("text", ""))
        if marker:
            add_risk_paragraph(doc, marker, text)
        else:
            doc.add_paragraph(text)
        return

    if block_type == "bullets":
        for item in block.get("items", []):
            if isinstance(item, dict):
                add_bullet(doc, str(item.get("text", "")), risk_marker(markers, item.get("risk")))
            else:
                add_bullet(doc, str(item))
        return

    if block_type == "table":
        headers = [str(header) for header in block.get("headers", [])]
        rows = normalize_table_rows(headers, block.get("rows", []))
        add_table(doc, headers, rows, red_markers=marker_values)
        return

    if block_type == "note":
        text = str(block.get("text", ""))
        if marker:
            add_risk_paragraph(doc, marker, text)
        else:
            paragraph = doc.add_paragraph()
            add_run(paragraph, text, bold=True)
        return

    if block_type == "warning":
        warning_marker = marker or markers.get(block.get("kind"), markers["missing"])
        add_warning(doc, warning_marker, str(block.get("text", "")))
        return

    if block_type == "image_placeholder":
        placeholder_marker = marker or markers["missing"]
        label = str(block.get("label") or block.get("text") or "Figure placeholder")
        add_placeholder_image(doc, placeholder_marker, label)
        return

    raise ValueError(f"Unsupported content block type: {block_type}")


def render_structured_section(doc, section: dict, markers: dict, marker_values: set[str]) -> None:
    title = str(section.get("title", "UNTITLED SECTION"))
    add_heading(doc, title)
    for block in section.get("blocks", []):
        render_block(doc, block, markers, marker_values)


def render_legacy_section(doc, section: str, config: dict, markers: dict, marker_values: set[str]) -> None:
    add_heading(doc, section)
    if section == "PIN DESCRIPTION":
        pins = config.get("required_pins", [])
        rows = [[pin, "TBD", markers["missing"]] for pin in pins] or [["TBD", "TBD", markers["missing"]]]
        add_table(doc, ["PIN", "NAME", "STATUS"], rows, red_markers=marker_values)
    elif section == "REVISION HISTORY":
        product = config.get("product", {})
        add_table(doc, ["REV", "DATE", "DESCRIPTION"], [[product.get("revision", ""), product.get("date", ""), "Initial draft."]], red_markers=marker_values)
    else:
        doc.add_paragraph(f"{section} content must be mapped from confirmed existing company product data and competitor reference material.")


def render_sections(doc, config: dict, markers: dict, marker_values: set[str]) -> None:
    sections = config.get("sections", [])
    for section in sections:
        if isinstance(section, dict):
            render_structured_section(doc, section, markers, marker_values)
        else:
            render_legacy_section(doc, str(section), config, markers, marker_values)


def input_documents(sources: dict) -> list[dict]:
    documents = sources.get("input_documents", [])
    if isinstance(documents, list):
        return [document for document in documents if isinstance(document, dict)]
    return []


def source_path_for_role(sources: dict, *roles: str) -> Path | None:
    for document in input_documents(sources):
        if document.get("role") in roles and document.get("path"):
            return Path(str(document["path"]))
    legacy_keys = {"style_template": "template"}
    for role in roles:
        key = legacy_keys.get(role, role)
        if sources.get(key):
            return Path(str(sources[key]))
    return None


def source_table_rows(sources: dict) -> list[list[str]]:
    rows: list[list[str]] = []
    for document in input_documents(sources):
        rows.append([
            str(document.get("role", "")),
            str(document.get("path", "")),
            str(document.get("purpose", "")),
        ])
    if rows:
        for key in ("analysis_report", "source_map", "asset_manifest"):
            if sources.get(key):
                rows.append([key, str(sources[key]), "generated support file"])
        return rows
    return [[role, str(path), "legacy source field"] for role, path in sources.items()]


def render_template_guidance(doc, config: dict, marker_values: set[str]) -> None:
    guidance = config.get("template_guidance") or []
    if not guidance:
        return
    add_heading(doc, "Template Guidance Disposition")
    rows = []
    for item in guidance:
        rows.append(
            [
                str(item.get("source_ref", "")),
                str(item.get("instruction", "")),
                str(item.get("decision", "")),
                str(item.get("evidence", "")),
            ]
        )
    add_table(doc, ["SOURCE REF", "INSTRUCTION", "DECISION", "EVIDENCE"], rows, red_markers=marker_values)


def include_audit_sections(config: dict) -> bool:
    options = config.get("output_options") or {}
    if isinstance(options, dict) and "include_audit_sections" in options:
        return bool(options["include_audit_sections"])
    return bool(config.get("include_audit_sections", False))


def local_name(element) -> str:
    return element.tag.rsplit("}", 1)[-1]


def attr_value(element, name: str) -> str | None:
    for key, value in element.attrib.items():
        if key.rsplit("}", 1)[-1] == name:
            return value
    return None


def header_footer_ref_count(sect_pr) -> int:
    return sum(1 for child in sect_pr if local_name(child) in {"headerReference", "footerReference"})


def has_two_columns(sect_pr) -> bool:
    for child in sect_pr:
        if local_name(child) == "cols" and attr_value(child, "num") == "2":
            return True
    return False


def section_properties(doc):
    return [element for element in doc._element.body.iter() if local_name(element) == "sectPr"]


def reusable_section_properties(doc):
    """Choose template section properties to keep after removing placeholder body."""
    sections = section_properties(doc)
    if not sections:
        return None
    for section in sections:
        if has_two_columns(section) and header_footer_ref_count(section):
            return deepcopy(section)
    for section in sections:
        if header_footer_ref_count(section):
            return deepcopy(section)
    for section in sections:
        if has_two_columns(section):
            return deepcopy(section)
    return deepcopy(sections[-1])


def clear_template_body(doc) -> None:
    """Remove template placeholder body while preserving reusable layout properties."""
    section = reusable_section_properties(doc)
    body = doc._element.body
    for child in list(body):
        body.remove(child)
    if section is not None:
        body.append(section)


def clear_story_part(part, text: str) -> None:
    try:
        part.is_linked_to_previous = False
    except Exception:
        pass
    element = part._element
    for child in list(element):
        element.remove(child)
    if text:
        part.add_paragraph(text)


def replace_headers_and_footers(doc, product: dict) -> None:
    name = str(product.get("name") or "<target_part_number>")
    subtitle = str(product.get("subtitle") or "Preliminary Semiconductor Datasheet")
    revision_date = f"{product.get('revision', '')} {product.get('date', '')}".strip()
    header_text = f"{name} | {subtitle}".strip(" |")
    footer_text = f"{name} {revision_date}".strip()
    for section in doc.sections:
        for story in (section.header, section.first_page_header, section.even_page_header):
            clear_story_part(story, header_text)
        for story in (section.footer, section.first_page_footer, section.even_page_footer):
            clear_story_part(story, footer_text)


def story_xml(tag: str, text: str) -> bytes:
    root = ET.Element(f"{{{W_NS['w']}}}{tag}")
    paragraph = ET.SubElement(root, f"{{{W_NS['w']}}}p")
    run = ET.SubElement(paragraph, f"{{{W_NS['w']}}}r")
    text_node = ET.SubElement(run, f"{{{W_NS['w']}}}t")
    text_node.text = text
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def rewrite_header_footer_parts(docx_path: Path, product: dict) -> None:
    name = str(product.get("name") or "<target_part_number>")
    subtitle = str(product.get("subtitle") or "Preliminary Semiconductor Datasheet")
    revision_date = f"{product.get('revision', '')} {product.get('date', '')}".strip()
    header_text = f"{name} | {subtitle}".strip(" |")
    footer_text = f"{name} {revision_date}".strip()
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".hdrftr.tmp")
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/header") and item.filename.endswith(".xml"):
                data = story_xml("hdr", header_text)
            elif item.filename.startswith("word/footer") and item.filename.endswith(".xml"):
                data = story_xml("ftr", footer_text)
            item.compress_type = zipfile.ZIP_DEFLATED
            target.writestr(item, data)
    tmp_path.replace(docx_path)


def is_annotation_part(name: str) -> bool:
    return name.startswith("word/comments") or name == "word/people.xml"


def remove_matching_children(root: ET.Element, predicate) -> bool:
    removed = False
    for parent in root.iter():
        for child in list(parent):
            if predicate(child):
                parent.remove(child)
                removed = True
    return removed


def strip_annotation_xml(name: str, data: bytes) -> bytes:
    if name == "[Content_Types].xml":
        root = ET.fromstring(data)
        changed = remove_matching_children(
            root,
            lambda child: (
                child.tag.endswith("}Override")
                and (
                    child.attrib.get("PartName", "").startswith("/word/comments")
                    or child.attrib.get("PartName") == "/word/people.xml"
                    or "comment" in child.attrib.get("ContentType", "").lower()
                )
            ),
        )
        return ET.tostring(root, encoding="utf-8", xml_declaration=True) if changed else data

    if name.endswith(".rels"):
        root = ET.fromstring(data)
        changed = remove_matching_children(
            root,
            lambda child: (
                child.tag.endswith("}Relationship")
                and (
                    "comment" in child.attrib.get("Type", "").lower()
                    or child.attrib.get("Type", "").endswith("/people")
                    or child.attrib.get("Target", "").startswith("comments")
                    or child.attrib.get("Target") == "people.xml"
                )
            ),
        )
        return ET.tostring(root, encoding="utf-8", xml_declaration=True) if changed else data

    if name.endswith(".xml"):
        root = ET.fromstring(data)
        annotation_tags = {"commentRangeStart", "commentRangeEnd", "commentReference", "annotationRef"}
        changed = remove_matching_children(
            root,
            lambda child: child.tag.rsplit("}", 1)[-1] in annotation_tags,
        )
        return ET.tostring(root, encoding="utf-8", xml_declaration=True) if changed else data

    return data


def strip_template_annotations(docx_path: Path) -> None:
    """Remove Word comment/annotation parts inherited from a source template."""
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            if is_annotation_part(item.filename):
                continue
            data = source.read(item.filename)
            if item.filename.endswith((".xml", ".rels")) or item.filename == "[Content_Types].xml":
                data = strip_annotation_xml(item.filename, data)
            item.compress_type = zipfile.ZIP_DEFLATED
            target.writestr(item, data)
    tmp_path.replace(docx_path)


def create_document(config: dict, output: Path) -> None:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required for DOCX generation. Install python-docx.") from exc

    sources = config.get("sources", {})
    template = source_path_for_role(sources, "style_template")
    can_clone_template = (
        template is not None
        and template.exists()
        and template.suffix.lower() in {".docx", ".docm", ".dotx", ".dotm"}
    )
    doc = Document(str(template)) if can_clone_template else Document()
    if can_clone_template:
        clear_template_body(doc)

    product = config.get("product", {})
    markers = config.get("risk_markers", SAMPLE_CONFIG["risk_markers"])
    marker_values = set(markers.values())
    name = product.get("name", "<target_part_number>")
    replace_headers_and_footers(doc, product)

    add_title(doc, name)
    doc.add_paragraph(product.get("subtitle", "Preliminary Semiconductor Datasheet"))
    doc.add_paragraph(f"{product.get('revision', '')}  {product.get('date', '')}".strip())
    doc.add_paragraph(product.get("compatibility", ""))

    if include_audit_sections(config):
        add_heading(doc, "Source and Risk Control")
        add_table(
            doc,
            ["SOURCE ROLE", "PATH", "PURPOSE"],
            source_table_rows(sources),
            red_markers=marker_values,
        )
        render_template_guidance(doc, config, marker_values)

        add_heading(doc, "Format Verification Requirements")
        for item in (
            "Render template and generated DOCX with render_document.py.",
            "Verify headings, lists, TOC, notes, tables, images, comments, headers, footers, and watermarks.",
            "Do not claim layout fidelity without render evidence.",
        ):
            doc.add_paragraph(item, style=None)

    render_sections(doc, config, markers, marker_values)

    for item in config.get("risk_items", []):
        marker = markers.get(item.get("kind"), markers["missing"])
        add_warning(doc, marker, item.get("text", "Unresolved item"))

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    strip_template_annotations(output)
    rewrite_header_footer_parts(output, product)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Generate a configurable preliminary datasheet DOCX.")
    parser.add_argument("--init-config", type=Path, help="Write a sample JSON config and exit")
    parser.add_argument("--config", type=Path, help="JSON config for generation")
    parser.add_argument("--output", type=Path, help="Output DOCX")
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    if args.init_config:
        init_config(args.init_config)
        print(f"Wrote sample config: {args.init_config}")
        return 0
    if not args.config or not args.output:
        print("ERROR: pass --config and --output, or use --init-config", file=sys.stderr)
        return 2
    try:
        create_document(load_config(args.config), args.output)
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    print(f"Generated datasheet: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
