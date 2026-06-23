"""Render DOCX/PDF documents to PNG pages for layout review."""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path


PDF_EXTENSIONS = {".pdf"}
WORD_EXTENSIONS = {".doc", ".docx", ".docm", ".dotx", ".dotm"}

INSTALL_HELP = """DOCX rendering requires one document-to-PDF engine:
- Best fidelity on Windows: Microsoft Word COM with pywin32.
- Fallback: LibreOffice 7.x+ with soffice on PATH, or set LIBREOFFICE_PATH.
- PDF rendering requires PyMuPDF (`fitz`).
- Diagnostic fallback: use `--engine docx-preview` for non-Word-fidelity layout previews when DOCX engines are unavailable.
"""


@dataclass
class RenderResult:
    input_path: Path
    pdf_path: Path
    png_paths: list[Path]
    engine: str
    dpi: int
    notes: list[str] | None = None


def import_fitz():
    try:
        import fitz  # type: ignore
    except Exception as exc:
        raise RuntimeError("PyMuPDF is required for PDF-to-PNG rendering. Install PyMuPDF.") from exc
    return fitz


def parse_pages(raw: str | None) -> list[int] | None:
    if not raw:
        return None
    pages: list[int] = []
    seen: set[int] = set()
    for part in raw.split(","):
        token = part.strip()
        if not token:
            continue
        if "-" in token:
            left, right = token.split("-", 1)
            start = int(left)
            end = int(right)
            if start < 1 or end < start:
                raise ValueError(f"invalid 1-based page range: {token}")
            values = range(start, end + 1)
        else:
            page = int(token)
            if page < 1:
                raise ValueError("page numbers are 1-based")
            values = [page]
        for page in values:
            if page not in seen:
                pages.append(page)
                seen.add(page)
    return pages


def safe_stem(path: Path) -> str:
    ascii_only = path.stem.encode("ascii", "ignore").decode("ascii")
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", ascii_only).strip("._-")
    return cleaned or "document"


def has_word_com() -> bool:
    if sys.platform != "win32":
        return False
    try:
        import win32com.client  # type: ignore  # noqa: F401
    except Exception:
        return False
    return True


def candidate_libreoffice_paths() -> list[str]:
    candidates: list[str] = []
    env_path = os.environ.get("LIBREOFFICE_PATH")
    if env_path:
        env_candidate = Path(env_path)
        if env_candidate.is_dir():
            executable = "soffice.exe" if sys.platform == "win32" else "soffice"
            candidates.append(str(env_candidate / executable))
        candidates.append(env_path)
    for executable in ("soffice", "libreoffice"):
        found = shutil.which(executable)
        if found:
            candidates.append(found)
    if sys.platform == "win32":
        candidates.extend(
            [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]
        )
    elif sys.platform == "darwin":
        candidates.append("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    else:
        candidates.extend(["/usr/bin/libreoffice", "/usr/local/bin/libreoffice", "/snap/bin/libreoffice"])
    unique: list[str] = []
    seen: set[str] = set()
    for candidate in candidates:
        if candidate and candidate not in seen:
            unique.append(candidate)
            seen.add(candidate)
    return unique


def find_libreoffice() -> str | None:
    for candidate in candidate_libreoffice_paths():
        if Path(candidate).exists():
            return candidate
    return None


def convert_with_word(input_path: Path, pdf_path: Path) -> Path:
    if sys.platform != "win32":
        raise RuntimeError("Microsoft Word COM is only available on Windows.")
    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore
    except Exception as exc:
        raise RuntimeError("pywin32 is required for Word COM rendering.") from exc

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(input_path.resolve()), ReadOnly=True, AddToRecentFiles=False)
        document.ExportAsFixedFormat(str(pdf_path.resolve()), 17)
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()
    return pdf_path


def find_converted_pdf(input_path: Path, out_dir: Path, before: set[Path], expected_pdf: Path) -> Path | None:
    expected_names = [
        expected_pdf.name,
        f"{input_path.stem}.pdf",
        f"{safe_stem(input_path)}.pdf",
    ]
    for name in expected_names:
        candidate = out_dir / name
        if candidate.exists():
            return candidate

    current = sorted(out_dir.glob("*.pdf"))
    new_pdfs = [path for path in current if path not in before]
    if len(new_pdfs) == 1:
        return new_pdfs[0]
    if new_pdfs:
        return max(new_pdfs, key=lambda path: path.stat().st_mtime)
    return None


def remove_stale_conversion_outputs(input_path: Path, out_dir: Path, expected_pdf: Path) -> None:
    for name in {expected_pdf.name, f"{input_path.stem}.pdf", f"{safe_stem(input_path)}.pdf"}:
        candidate = out_dir / name
        if candidate.exists():
            candidate.unlink()


def convert_with_libreoffice(input_path: Path, pdf_path: Path) -> Path:
    executable = find_libreoffice()
    if not executable:
        searched = "\n".join(f"- {candidate}" for candidate in candidate_libreoffice_paths())
        raise RuntimeError("LibreOffice executable not found.\n" + INSTALL_HELP + "\nSearched:\n" + searched)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    remove_stale_conversion_outputs(input_path, pdf_path.parent, pdf_path)
    before = set(pdf_path.parent.glob("*.pdf"))
    with tempfile.TemporaryDirectory(prefix=".lo-profile-", dir=pdf_path.parent) as profile_raw, tempfile.TemporaryDirectory(
        prefix=".lo-home-", dir=pdf_path.parent
    ) as home_raw, tempfile.TemporaryDirectory(prefix=".lo-runtime-", dir=pdf_path.parent) as runtime_raw:
        profile_dir = Path(profile_raw)
        home_dir = Path(home_raw)
        runtime_dir = Path(runtime_raw)
        runtime_dir.chmod(0o700)
        env = os.environ.copy()
        env["HOME"] = str(home_dir)
        env["XDG_RUNTIME_DIR"] = str(runtime_dir)
        env["SAL_USE_VCLPLUGIN"] = env.get("SAL_USE_VCLPLUGIN", "svp")
        env.pop("DISPLAY", None)
        cmd = [
            executable,
            f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--nodefault",
            "--nolockcheck",
            "--invisible",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(input_path.resolve()),
        ]
        completed = subprocess.run(cmd, capture_output=True, text=True, timeout=120, env=env)
    if completed.returncode != 0:
        raise RuntimeError(f"LibreOffice failed:\nSTDOUT:\n{completed.stdout}\nSTDERR:\n{completed.stderr}")
    produced = find_converted_pdf(input_path, pdf_path.parent, before, pdf_path)
    if produced is None:
        produced_files = "\n".join(f"- {path.name}" for path in sorted(pdf_path.parent.iterdir()))
        raise RuntimeError(
            "LibreOffice did not produce a discoverable PDF.\n"
            f"Expected: {pdf_path.name}, {input_path.stem}.pdf, or {safe_stem(input_path)}.pdf\n"
            f"Output directory files:\n{produced_files}\n"
            f"STDOUT:\n{completed.stdout}\nSTDERR:\n{completed.stderr}"
        )
    if produced.resolve() != pdf_path.resolve():
        if pdf_path.exists():
            pdf_path.unlink()
        produced.replace(pdf_path)
    return pdf_path


def convert_to_pdf(input_path: Path, pdf_path: Path, engine: str) -> tuple[Path, str]:
    suffix = input_path.suffix.lower()
    if suffix in PDF_EXTENSIONS:
        return input_path, "pdf"
    if suffix not in WORD_EXTENSIONS:
        raise ValueError(f"unsupported input type: {input_path.suffix}")
    if engine == "word":
        return convert_with_word(input_path, pdf_path), "word"
    if engine == "libreoffice":
        return convert_with_libreoffice(input_path, pdf_path), "libreoffice"
    if engine != "auto":
        raise ValueError(f"unknown engine: {engine}")
    errors: list[str] = []
    if has_word_com():
        try:
            return convert_with_word(input_path, pdf_path), "word"
        except Exception as exc:
            errors.append(f"Word COM failed: {exc}")
    else:
        errors.append("Word COM unavailable.")
    if find_libreoffice():
        try:
            return convert_with_libreoffice(input_path, pdf_path), "libreoffice"
        except Exception as exc:
            errors.append(f"LibreOffice failed: {exc}")
    else:
        errors.append("LibreOffice unavailable.")
    raise RuntimeError("No document-to-PDF engine succeeded.\n" + "\n".join(errors) + "\n" + INSTALL_HELP)


def length_to_points(value, default: float) -> float:
    try:
        return float(value.pt)
    except Exception:
        return default


def docx_text_of(node) -> str:
    return "".join(text.text or "" for text in node.findall(".//w:t", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}))


def local_name(node) -> str:
    return node.tag.rsplit("}", 1)[-1]


def docx_column_count(document) -> int:
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for section in document._element.body.iter():
        if local_name(section) != "sectPr":
            continue
        cols = section.find("./w:cols", namespace)
        if cols is not None:
            raw = cols.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num")
            if raw and raw.isdigit():
                return max(1, int(raw))
    return 1


def draw_textbox(page, rect, text: str, *, size: float = 8.5, color=(0, 0, 0)) -> None:
    page.insert_textbox(rect, text[:1200], fontsize=size, fontname="helv", color=color, align=0)


def render_docx_preview_pdf(input_path: Path, pdf_path: Path) -> Path:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required for docx-preview rendering.") from exc
    fitz = import_fitz()

    document = Document(str(input_path))
    section = document.sections[0]
    page_width = length_to_points(section.page_width, 595.0)
    page_height = length_to_points(section.page_height, 842.0)
    left = length_to_points(section.left_margin, 54.0)
    right = length_to_points(section.right_margin, 54.0)
    top = length_to_points(section.top_margin, 54.0)
    bottom = length_to_points(section.bottom_margin, 54.0)
    columns = docx_column_count(document)
    gap = 18.0 if columns > 1 else 0.0
    usable_width = page_width - left - right
    column_width = (usable_width - gap * (columns - 1)) / columns
    header_text = " | ".join(p.text for p in section.header.paragraphs if p.text.strip())[:180]
    footer_text = " | ".join(p.text for p in section.footer.paragraphs if p.text.strip())[:180]

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    preview = fitz.open()
    page = None
    column = 0
    y = top

    def add_page():
        new_page = preview.new_page(width=page_width, height=page_height)
        new_page.draw_rect(fitz.Rect(0, 0, page_width, page_height), color=(0.78, 0.78, 0.78), width=0.5)
        if header_text:
            draw_textbox(new_page, fitz.Rect(left, 20, page_width - right, top - 8), header_text, size=7.5, color=(0.55, 0, 0))
        if footer_text:
            draw_textbox(new_page, fitz.Rect(left, page_height - bottom + 10, page_width - right, page_height - 18), footer_text, size=7.5, color=(0.35, 0.35, 0.35))
        if columns > 1:
            for index in range(1, columns):
                x = left + index * column_width + (index - 0.5) * gap
                new_page.draw_line((x, top), (x, page_height - bottom), color=(0.86, 0.86, 0.86), width=0.5)
        return new_page

    def current_x() -> float:
        return left + column * (column_width + gap)

    def advance(height: float) -> None:
        nonlocal page, column, y
        if page is None:
            page = add_page()
        if y + height <= page_height - bottom:
            return
        if column + 1 < columns:
            column += 1
            y = top
            return
        page = add_page()
        column = 0
        y = top

    def draw_paragraph(text: str, style: str | None, drawing_count: int) -> None:
        nonlocal page, y
        clean = " ".join(text.split())
        if not clean and drawing_count == 0:
            return
        is_heading = bool(style and ("Heading" in style or "标题" in style or "Title" in style))
        font_size = 11.5 if is_heading else 8.2
        line_height = font_size + 3.0
        estimated_lines = max(1, min(10, (len(clean) // max(24, int(column_width / (font_size * 0.45)))) + 1))
        box_height = estimated_lines * line_height + (8 if is_heading else 4) + drawing_count * 40
        advance(box_height)
        if page is None:
            page = add_page()
        rect = fitz.Rect(current_x(), y, current_x() + column_width, y + box_height)
        if is_heading:
            page.draw_rect(rect, color=(0.75, 0.05, 0.05), fill=(1, 0.96, 0.96), width=0.6)
            draw_textbox(page, rect + (3, 2, -3, -2), clean.upper(), size=font_size, color=(0.35, 0, 0))
        else:
            draw_textbox(page, rect, clean, size=font_size)
        for index in range(drawing_count):
            top_y = y + estimated_lines * line_height + index * 40
            image_rect = fitz.Rect(current_x(), top_y, current_x() + min(column_width, 130), top_y + 34)
            page.draw_rect(image_rect, color=(0.2, 0.45, 0.75), fill=(0.88, 0.93, 0.98), width=0.6)
            draw_textbox(page, image_rect + (3, 8, -3, -3), "FIGURE / IMAGE", size=7.5, color=(0.1, 0.25, 0.45))
        y += box_height + 4

    def draw_table(text: str, row_count: int) -> None:
        nonlocal page, y
        rows = max(1, min(row_count, 12))
        height = 17 + rows * 13
        advance(height)
        if page is None:
            page = add_page()
        x0 = current_x()
        rect = fitz.Rect(x0, y, x0 + column_width, y + height)
        page.draw_rect(rect, color=(0.2, 0.2, 0.2), fill=(0.97, 0.97, 0.97), width=0.6)
        page.draw_rect(fitz.Rect(x0, y, x0 + column_width, y + 15), color=(0.2, 0.2, 0.2), fill=(0.86, 0.9, 0.94), width=0.6)
        for row in range(1, rows + 1):
            yy = y + 15 + row * 13
            page.draw_line((x0, yy), (x0 + column_width, yy), color=(0.72, 0.72, 0.72), width=0.4)
        page.draw_line((x0 + column_width / 2, y), (x0 + column_width / 2, y + height), color=(0.72, 0.72, 0.72), width=0.4)
        label = "TABLE"
        if text:
            label += ": " + " ".join(text.split())[:120]
        draw_textbox(page, fitz.Rect(x0 + 3, y + 2, x0 + column_width - 3, y + height - 2), label, size=7.2)
        y += height + 6

    for child in list(document._element.body):
        name = local_name(child)
        if name == "p":
            style = None
            paragraph_id = child.find(".//w:pStyle", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            if paragraph_id is not None:
                style = paragraph_id.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
            drawing_count = len(child.findall(".//w:drawing", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}))
            draw_paragraph(docx_text_of(child), style, drawing_count)
        elif name == "tbl":
            row_count = len(child.findall(".//w:tr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}))
            draw_table(docx_text_of(child), row_count)

    if page is None:
        add_page()
    preview.save(pdf_path)
    preview.close()
    return pdf_path


def render_pdf_to_pngs(pdf_path: Path, out_dir: Path, pages: list[int] | None, dpi: int) -> list[Path]:
    fitz = import_fitz()
    out_dir.mkdir(parents=True, exist_ok=True)
    png_paths: list[Path] = []
    matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0)
    with fitz.open(pdf_path) as pdf:
        selected = pages or list(range(1, pdf.page_count + 1))
        invalid = [page for page in selected if page > pdf.page_count]
        if invalid:
            raise ValueError(f"requested pages exceed PDF page count {pdf.page_count}: {invalid}")
        for page_number in selected:
            pix = pdf[page_number - 1].get_pixmap(matrix=matrix, alpha=False)
            output = out_dir / f"page-{page_number:03d}.png"
            pix.save(output)
            png_paths.append(output)
    return png_paths


def render_document(input_path: Path, out_dir: Path, engine: str, pages: str | None, dpi: int) -> RenderResult:
    if not input_path.exists():
        raise FileNotFoundError(input_path)
    document_out_dir = out_dir / safe_stem(input_path)
    pdf_path = document_out_dir / f"{safe_stem(input_path)}.pdf"
    if engine == "docx-preview":
        if input_path.suffix.lower() not in WORD_EXTENSIONS:
            raise ValueError("docx-preview is only available for Word documents")
        converted_pdf = render_docx_preview_pdf(input_path, pdf_path)
        png_paths = render_pdf_to_pngs(converted_pdf, document_out_dir, parse_pages(pages), dpi)
        result = RenderResult(
            input_path,
            converted_pdf,
            png_paths,
            "docx-preview",
            dpi,
            ["docx-preview is not a Word-fidelity render; use it only for gross layout comparison when DOCX engines are unavailable."],
        )
        write_manifest(result, document_out_dir / "render_manifest.json")
        return result
    converted_pdf, actual_engine = convert_to_pdf(input_path, pdf_path, engine)
    png_paths = render_pdf_to_pngs(converted_pdf, document_out_dir, parse_pages(pages), dpi)
    result = RenderResult(input_path, converted_pdf, png_paths, actual_engine, dpi)
    write_manifest(result, document_out_dir / "render_manifest.json")
    return result


def write_manifest(result: RenderResult, manifest_path: Path) -> None:
    data = {
        "input": str(result.input_path.resolve()),
        "pdf": str(result.pdf_path.resolve()),
        "engine": result.engine,
        "dpi": result.dpi,
        "pngs": [str(path.resolve()) for path in result.png_paths],
    }
    if result.notes:
        data["notes"] = result.notes
    manifest_path.write_text(json.dumps(data, indent=2), encoding="utf-8")


def probe() -> int:
    print(f"Python: {sys.version.split()[0]}")
    try:
        fitz = import_fitz()
        print(f"PyMuPDF: {fitz.VersionBind}")
    except Exception as exc:
        print(f"PyMuPDF: unavailable ({exc})")
    print(f"Word COM available: {'yes' if has_word_com() else 'no'}")
    print(f"LibreOffice available: {find_libreoffice() or 'no'}")
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Render DOCX/PDF files to PNG pages for visual layout review.")
    parser.add_argument("input", nargs="?", type=Path, help="Input .docx/.doc/.pdf file")
    parser.add_argument("--out-dir", type=Path, default=Path("tmp") / "rendered")
    parser.add_argument("--engine", choices=["auto", "word", "libreoffice", "docx-preview"], default="auto")
    parser.add_argument("--pages", help="1-based pages or ranges, for example: 1,3-5")
    parser.add_argument("--dpi", type=int, default=160)
    parser.add_argument("--probe", action="store_true")
    parser.add_argument("--install-help", action="store_true")
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    if args.install_help:
        print(INSTALL_HELP.strip())
        return 0
    if args.probe:
        return probe()
    if not args.input:
        print("ERROR: input is required unless --probe is used", file=sys.stderr)
        return 2
    try:
        result = render_document(args.input, args.out_dir, args.engine, args.pages, args.dpi)
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    print(f"Rendered with engine: {result.engine}")
    print(f"PDF: {result.pdf_path}")
    for path in result.png_paths:
        print(f"PNG: {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
